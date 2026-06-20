"""
Nyaya AI -- Project Report to DOCX Converter
Converts PROJECT_REPORT.md into a professionally formatted Word document.

Usage:
    python scripts/convert_report_to_docx.py

Output:
    NyayaAI_Project_Report.docx (in project root)
"""

import re
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Try UTF-8 console output
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# --- Paths ------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).parent.parent
MARKDOWN_PATH = PROJECT_ROOT / "PROJECT_REPORT.md"
OUTPUT_PATH = PROJECT_ROOT / "NyayaAI_Project_Report.docx"
SCREENSHOTS_DIR = PROJECT_ROOT / "screenshots"

# --- Colours (Nyaya brand) --------------------------------------------------
INK = RGBColor(0x1A, 0x2E, 0x44)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
FOREST = RGBColor(0x2D, 0x4B, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF5, 0xF2, 0xED)
GOLD_HEX = "B8860B"
INK_HEX = "1A2E44"


def set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    shading_elm = parse_xml(
        '<w:shd %s w:fill="%s" w:val="clear"/>' % (nsdecls("w"), color_hex)
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_hyperlink(paragraph, text: str, url: str):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = parse_xml(
        '<w:hyperlink %s r:id="%s" %s>'
        '  <w:r><w:rPr><w:color w:val="%s"/><w:u w:val="single"/></w:rPr>'
        "  <w:t>%s</w:t></w:r>"
        "</w:hyperlink>"
        % (nsdecls("w"), r_id, nsdecls("r"), GOLD_HEX, text)
    )
    paragraph._p.append(hyperlink)


def parse_markdown_to_docx(doc: Document, md_text: str):
    """Parse the markdown text and build the DOCX content."""

    # --- Title Page ---------------------------------------------------------
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Nyaya AI")
    run.font.size = Pt(36)
    run.font.color.rgb = GOLD
    run.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("AI-Powered Indian Legal Research Platform")
    run.font.size = Pt(18)
    run.font.color.rgb = INK

    doc.add_paragraph()

    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline_p.add_run("Project Report")
    run.font.size = Pt(16)
    run.font.color.rgb = FOREST

    for _ in range(4):
        doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_p.add_run(
        "Submitted by: Nyaya AI Development Team\nDate: June 2026\nVersion: 1.0.0"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = MED_GRAY

    doc.add_page_break()

    # --- Parse body ---------------------------------------------------------
    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []
    table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- Code blocks ----------------------------------------------------
        if stripped.startswith("```"):
            if in_code_block:
                doc.add_paragraph("")
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                p.style = doc.styles["Normal"]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
                run.font.color.rgb = DARK_GRAY
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(stripped)
            i += 1
            continue

        # --- Tables ---------------------------------------------------------
        if "|" in stripped and stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            if table_buffer:
                create_table(doc, table_buffer)
                table_buffer = []
                doc.add_paragraph()

        # --- Headings -------------------------------------------------------
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            add_heading(doc, text, 1)
            i += 1
            continue
        elif stripped.startswith("### "):
            text = stripped[4:].strip()
            add_heading(doc, text, 2)
            i += 1
            continue
        elif stripped.startswith("#### "):
            text = stripped[5:].strip()
            add_heading(doc, text, 3)
            i += 1
            continue

        # --- Screenshots placeholder: replace note with actual images ------
        if "> **Note:** Screenshots are not available" in stripped:
            insert_screenshots(doc)
            i += 1
            continue

        # --- Figure captions ------------------------------------------------
        if stripped.startswith("**Figure") and stripped.endswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped.replace("**", ""))
            run.bold = True
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = MED_GRAY
            i += 1
            continue

        # --- Empty lines ----------------------------------------------------
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # --- Bullet points --------------------------------------------------
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            add_bullet(doc, text)
            i += 1
            continue

        # --- Numbered lists -------------------------------------------------
        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            add_bullet(doc, text, numbered=True)
            i += 1
            continue

        # --- Regular paragraphs ---------------------------------------------
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(15)
        parse_inline_text(p, stripped)
        i += 1

    if table_buffer:
        create_table(doc, table_buffer)


def add_heading(doc, text: str, level: int):
    """Add a styled heading."""
    p = doc.add_paragraph()

    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.size = Pt(20)
        run.font.color.rgb = GOLD
        run.bold = True
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr %s>'
            '  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="B8860B"/>'
            '</w:pBdr>' % nsdecls("w")
        )
        pPr.append(pBdr)

    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.color.rgb = INK
        run.bold = True

    elif level == 3:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.color.rgb = FOREST
        run.bold = True


_list_counter = [0]


def add_bullet(doc, text: str, numbered: bool = False):
    """Add a bullet or numbered list item."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27)

    if not numbered:
        run = p.add_run("- ")
    else:
        _list_counter[0] += 1
        run = p.add_run(str(_list_counter[0]) + ". ")

    run.font.size = Pt(11)
    run.font.color.rgb = INK
    parse_inline_text(p, text)


def parse_inline_text(p, text: str):
    """Parse inline **bold**, `code`, and [link](url)."""
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = INK
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = FOREST
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            idx = part.index("](")
            link_text = part[1:idx]
            link_url = part[idx + 2 : -1]
            add_hyperlink(p, link_text, link_url)
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY


# --- Screenshot configuration -------------------------------------------------
SCREENSHOTS = [
    {
        "filename": "01_search_results.png",
        "figure": "Figure 7.1",
        "caption": "Search Results Page -- Keyword search for 'Maneka Gandhi' showing ranked results with case titles, court badges, dates, and citation numbers. High-confidence case name matches are ranked first.",
        "width": Cm(15),
    },
    {
        "filename": "02_case_detail.png",
        "figure": "Figure 7.2",
        "caption": "Case Detail Page -- Full judgment text for Maneka Gandhi v. Union of India displayed in serif reading view with metadata header (court, date, citation, bench) showing the judgment opening.",
        "width": Cm(15),
    },
    {
        "filename": "03_bare_acts.png",
        "figure": "Figure 7.3",
        "caption": "Bare Acts Reader -- Bharatiya Nyaya Sanhita, 2023 showing the act title, year, table of contents with expandable chapters in the sidebar, and section text in the main content area.",
        "width": Cm(15),
    },
    {
        "filename": "04_citation_graph.png",
        "figure": "Figure 7.4",
        "caption": "Citation Graph -- Interactive force-directed graph for Maneka Gandhi v. Union of India showing the root case (gold), cases it cites (blue), and subsequent cases that cited it (green). Nodes are clickable for navigation.",
        "width": Cm(15),
    },
    {
        "filename": "05_dictionary.png",
        "figure": "Figure 7.5",
        "caption": "Legal Dictionary Page -- Browseable list of legal terms and Latin maxims with AI-powered explanations. Each term card shows the term name and a preview of its legal definition.",
        "width": Cm(15),
    },
    {
        "filename": "06_rights.png",
        "figure": "Figure 7.6",
        "caption": "Know Your Rights Page -- Situation-based fundamental rights cards showing plain English explanations of rights like Right to Privacy under Article 21, with actionable guidance and landmark case references.",
        "width": Cm(15),
    },
    {
        "filename": "07_judges.png",
        "figure": "Figure 7.7",
        "caption": "Judge Profiles Page -- AI-powered judicial analytics showing judge profiles with ideological tendency scores, practice area breakdowns, and lists of recent judgments for Supreme Court judges.",
        "width": Cm(15),
    },
    {
        "filename": "08_maxims.png",
        "figure": "Figure 7.8",
        "caption": "Legal Maxims Page -- Browseable collection of Latin legal maxims with AI-powered explanations. Each maxim card shows the Latin phrase and its legal meaning with case examples.",
        "width": Cm(15),
    },
]


def insert_screenshots(doc):
    """Insert all screenshots into the document after Chapter 7 heading."""
    if not SCREENSHOTS_DIR.exists():
        print("[WARN] Screenshots directory not found: " + str(SCREENSHOTS_DIR))
        return

    for shot in SCREENSHOTS:
        img_path = SCREENSHOTS_DIR / shot["filename"]
        if not img_path.exists():
            print("[WARN] Screenshot not found: " + shot["filename"])
            continue

        # Add screenshot image, centered
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run()
        run.add_picture(str(img_path), width=shot["width"])

        # Add figure number (bold + italic)
        fig_p = doc.add_paragraph()
        fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig_p.paragraph_format.space_after = Pt(2)
        run = fig_p.add_run(shot["figure"])
        run.bold = True
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = MED_GRAY

        # Add caption
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(14)
        run = cap_p.add_run(shot["caption"])
        run.font.size = Pt(9)
        run.font.color.rgb = MED_GRAY
        run.italic = True


def create_table(doc, table_lines):
    """Create a formatted table from markdown table lines."""
    header_cells = [c.strip() for c in table_lines[0].split("|") if c.strip()]
    data_rows = []
    for row_line in table_lines[2:]:
        cells = [c.strip() for c in row_line.split("|") if c.strip()]
        if cells:
            data_rows.append(cells)

    if not data_rows:
        return

    num_cols = len(header_cells)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, INK_HEX)

    for i, row_cells in enumerate(data_rows):
        for j, cell_text in enumerate(row_cells):
            if j >= num_cols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(8.5)
            run.font.color.rgb = DARK_GRAY
            if i % 2 == 0:
                set_cell_shading(cell, "F5F2ED")

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)


def main():
    print("=" * 60)
    print("  Nyaya AI -- Report to DOCX Converter")
    print("=" * 60)

    if not MARKDOWN_PATH.exists():
        print("[ERR] Markdown file not found: " + str(MARKDOWN_PATH))
        print("   Run: python scripts/convert_report_to_docx.py")
        return

    md_text = MARKDOWN_PATH.read_text(encoding="utf-8")
    print("[OK] Read %d characters from %s" % (len(md_text), MARKDOWN_PATH.name))

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    parse_markdown_to_docx(doc, md_text)

    # Add footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = footer_p.add_run("-- Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MED_GRAY

    fld_char_begin = parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls("w"))
    instr_text = parse_xml(
        '<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls("w")
    )
    fld_char_end = parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls("w"))

    run2 = footer_p.add_run()
    run2._r.append(fld_char_begin)
    run2._r.append(instr_text)
    run2._r.append(fld_char_end)

    run3 = footer_p.add_run(" --")
    run3.font.size = Pt(9)
    run3.font.color.rgb = MED_GRAY

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    file_size = OUTPUT_PATH.stat().st_size / 1024
    print("[OK] Saved: %s (%.0f KB)" % (OUTPUT_PATH.name, file_size))
    print("=" * 60)


if __name__ == "__main__":
    main()
